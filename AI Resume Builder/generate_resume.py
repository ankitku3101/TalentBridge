from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def generate_resume_from_template(data, output_file):

    doc = Document()

    # Add profile picture to the resume (optional)
    if data.get("profile_picture"):
        doc.add_picture(data["profile_picture"], width=Inches(1.5))
        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center the image
        doc.add_paragraph()

    # Title Section: Name and Job Title
    title_section = doc.add_heading(level=1)
    name_run = title_section.add_run(f"{data['name']}\n")
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(54, 95, 145)  # Stylish blue color
    title_section.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    job_title = doc.add_paragraph(data['job_title'])
    job_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    job_title_run = job_title.runs[0]
    job_title_run.italic = True
    job_title_run.font.size = Pt(14)
    job_title_run.font.color.rgb = RGBColor(0, 102, 102)  # Subtle green color

    doc.add_paragraph()  # Add spacing

    # Contact Information
    doc.add_heading("Contact Information", level=2).font.color.rgb = RGBColor(0, 0, 0)
    contact_info = f"Address: {data['address']}\nPhone: {data['phone']}\nEmail: {data['email']}"
    contact_para = doc.add_paragraph(contact_info)
    contact_para.paragraph_format.space_after = Pt(10)
    doc.add_paragraph()

    # Career Objective
    career_heading = doc.add_heading("Career Objective", level=2)
    career_heading.font.color.rgb = RGBColor(54, 95, 145)
    career_para = doc.add_paragraph(data['career_objective'])
    career_para.style.font.size = Pt(12)
    career_para.style.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()

    # Educational Qualifications
    edu_heading = doc.add_heading("Educational Qualifications", level=2)
    edu_heading.font.color.rgb = RGBColor(54, 95, 145)
    for edu in data['education']:
        para = doc.add_paragraph(style="List Bullet")
        degree_run = para.add_run(f"{edu['degree']} - {edu['institute']}\n")
        degree_run.bold = True
        degree_run.font.size = Pt(12)
        details_run = para.add_run(f"Year: {edu['year']} - CGPA: {edu.get('cgpa', 'N/A')}")
        details_run.font.size = Pt(11)
    doc.add_paragraph()

    # Technical Skills
    skills_heading = doc.add_heading("Technical Skills", level=2)
    skills_heading.font.color.rgb = RGBColor(54, 95, 145)
    for skill in data['skills']:
        skill_para = doc.add_paragraph(skill, style="List Bullet")
        skill_para.style.font.size = Pt(11)
    doc.add_paragraph()

    # Projects Undertaken
    projects_heading = doc.add_heading("Projects Undertaken", level=2)
    projects_heading.font.color.rgb = RGBColor(54, 95, 145)
    for project in data['projects']:
        para = doc.add_paragraph(style="List Bullet")
        project_name = para.add_run(f"{project['name']}\n")
        project_name.bold = True
        project_name.font.size = Pt(12)
        project_desc = para.add_run(f"{project['description']}")
        project_desc.font.size = Pt(11)
    doc.add_paragraph()

    # Footer
    footer_section = doc.sections[-1].footer
    footer_para = footer_section.paragraphs[0]
    footer_para.text = "Generated using AI Resume Builder"
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_para.runs[0].font.size = Pt(10)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)  # Gray color

    # Save the document
    doc.save(output_file)